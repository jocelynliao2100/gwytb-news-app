import streamlit as st
from docx import Document
import pandas as pd
import jieba.analyse

st.set_page_config(page_title="國台辦新聞稿分析", layout="wide")

st.title("🇨🇳 國台辦新聞稿分析首頁")

st.markdown("""
歡迎使用 **國台辦新聞稿分析系統**！

請先在左側選單中選擇你要分析的功能，並上傳對應的 Word 檔案（.docx）。

本系統提供以下三個分析模組：

1. 📂 **五大欄目基本資訊**：分析新聞數量、時間分布與表格呈現  
2. 🔍 **關鍵字分析**：抽取新聞標題中的高頻關鍵詞並視覺化  
3. 🌐 **「交往交流」欄目分析**：聚焦與「青年」、「台胞」有關的欄目新聞

---

""")

# 側邊欄選單
menu = st.sidebar.radio("📁 選擇分析模組", [
    "首頁",
    "五大欄目基本資訊",
    "關鍵字分析",
    "「交往交流」欄目分析"
])

# 共用的檔案處理函式
def parse_docx(file):
    doc = Document(file)
    data = []
    for para in doc.paragraphs:
        line = para.text.strip()
        if line.startswith("[") and "]" in line:
            date = line.split("]")[0].strip("[]")
            title = line.split("]")[-1].strip()
            if title:
                data.append({"日期": date, "標題": title})
    return pd.DataFrame(data)

# 僅非首頁時顯示上傳
if menu != "首頁":
    uploaded_file = st.file_uploader("請上傳 Word 檔（例如：國台辦原始碼）", type="docx")

    if uploaded_file:
        df = parse_docx(uploaded_file)

        if menu == "五大欄目基本資訊":
            st.subheader("📂 五大欄目基本資訊")
            st.write(f"共載入 {len(df)} 筆新聞")
            st.dataframe(df)
            with st.expander("📊 日期分布圖"):
                st.bar_chart(df["日期"].value_counts().sort_index())

        elif menu == "關鍵字分析":
            st.subheader("🔍 關鍵字分析")
            full_text = " ".join(df["標題"])
            top_keywords = jieba.analyse.extract_tags(full_text, topK=20, withWeight=True)
            keyword_df = pd.DataFrame(top_keywords, columns=["關鍵字", "權重"])
            st.dataframe(keyword_df)
            st.bar_chart(keyword_df.set_index("關鍵字"))

        elif menu == "「交往交流」欄目分析":
            st.subheader("🌐 「交往交流」欄目分析")
            filtered_df = df[df["標題"].str.contains("交流|交往|台胞|台青|座談|研習|文創|聯誼")]
            st.write(f"共找到 {len(filtered_df)} 筆相關新聞")
            st.dataframe(filtered_df)
            with st.expander("📆 日期統計圖"):
                st.bar_chart(filtered_df["日期"].value_counts().sort_index())
    else:
        st.info("請先上傳檔案以啟用此功能")


import streamlit as st
from docx import Document
from bs4 import BeautifulSoup
import pandas as pd
import re

st.subheader("📂 五大欄目基本資訊（多檔分析）")

uploaded_files = st.file_uploader("請上傳多個國台辦 Word 檔案（HTML 原始碼形式）", type="docx", accept_multiple_files=True)

def extract_news_from_docx(file):
    doc = Document(file)
    html = "\n".join(p.text for p in doc.paragraphs)
    soup = BeautifulSoup(html, "html.parser")

    # 欄目名稱來自 <meta name="ColumnName" content="台办动态" />
    column_meta = soup.find("meta", attrs={"name": "ColumnName"})
    column_name = column_meta["content"] if column_meta else "未知欄目"

    news_items = []
    for li in soup.find_all("li"):
        date_tag = li.find("span")
        link_tag = li.find("a")
        if date_tag and link_tag:
            date = date_tag.text.strip("[] ")
            title = link_tag.get("title", "").strip()
            if re.match(r"\d{4}-\d{2}-\d{2}", date) and title:
                news_items.append({
                    "日期": date,
                    "標題": title,
                    "欄目": column_name
                })
    return news_items

if uploaded_files:
    all_data = []
    for f in uploaded_files:
        all_data.extend(extract_news_from_docx(f))

    df = pd.DataFrame(all_data)
    st.success(f"✅ 共解析出 {len(df)} 則新聞，涵蓋 {df['欄目'].nunique()} 個欄目")
    st.dataframe(df)

    st.markdown("### 🧭 各欄目新聞數量")
    st.bar_chart(df["欄目"].value_counts())

    st.markdown("### 📅 各欄目按日期分布")
    pivot = df.groupby(["欄目", "日期"]).size().reset_index(name="數量")
    pivot_wide = pivot.pivot(index="日期", columns="欄目", values="數量").fillna(0)
    st.line_chart(pivot_wide)

else:
    st.info("請上傳至少一個 .docx 檔案以進行分析")
